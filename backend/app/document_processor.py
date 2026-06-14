from docx import Document


class DocumentProcessor:
    """Handle DOCX document processing"""

    def apply_text_corrections(
        self,
        input_path: str,
        output_path: str,
        text_corrections: list,
    ) -> int:
        """
        Áp dụng danh sách sửa lỗi dựa trên nội dung text (không cần element ID).
        Dùng cho kết quả từ RAG vì RAG trả về original_text, không trả về vị trí cấu trúc.

        Args:
            text_corrections: list of {'original_text': str, 'new_text': str}

        Returns:
            số lượng thay thế thực sự được thực hiện
        """
        doc = Document(input_path)
        applied = 0

        for corr in text_corrections:
            old = (corr.get('original_text') or '').strip()
            new = (corr.get('new_text') or '').strip()
            if not old or not new or old == new:
                continue

            # Duyệt paragraphs cấp cao nhất
            for para in doc.paragraphs:
                if old in para.text:
                    if self._replace_in_paragraph(para, old, new):
                        applied += 1

            # Duyệt tất cả các cell trong tất cả bảng
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if old in para.text:
                                if self._replace_in_paragraph(para, old, new):
                                    applied += 1

        doc.save(output_path)
        return applied

    def _replace_in_paragraph(self, para, old_text: str, new_text: str) -> bool:
        """
        Thay old_text bằng new_text trong paragraph.
        Thử thay trong từng run đơn (giữ định dạng), fallback sang merge toàn bộ runs.
        """
        # Ưu tiên: nếu old_text nằm gọn trong một run, chỉ sửa run đó
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text, 1)
                return True

        # Fallback: text rải rác nhiều runs — merge về run đầu, xoá phần còn lại
        full_text = para.text
        if old_text in full_text:
            replaced = full_text.replace(old_text, new_text, 1)
            if para.runs:
                para.runs[0].text = replaced
                for run in para.runs[1:]:
                    run.text = ''
            else:
                para.add_run(replaced)
            return True

        return False
