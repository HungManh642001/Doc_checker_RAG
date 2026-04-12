"""
Script để tạo tài liệu mẫu cho testing
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os


def create_sample_main_document():
    """Tạo tài liệu mẫu cần thẩm định"""
    doc = Document()
    
    # Title
    title = doc.add_heading('HỢP ĐỒNG MUA BÁN HÀNG HÓA', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Date
    date_para = doc.add_paragraph()
    date_para.add_run('Ngày lập: 08/04/2026').bold = True
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Content
    doc.add_heading('I. THÔNG TIN BÊN GIAO DỊCH', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Bên mua (A): ').bold = True
    p.add_run('CÔNG TY TNHH ABC')  # Inconsistent format - should be "Công ty"
    
    p = doc.add_paragraph()
    p.add_run('Đại diện: ').bold = True
    p.add_run('Ông Nguyễn Văn A')
    
    p = doc.add_paragraph()
    p.add_run('Bên bán (B): ').bold = True
    p.add_run('Công ty XYZ')
    
    # Add unit notation errors for RAG testing
    doc.add_heading('II. THÔNG SỐ KỸ THUẬT', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Kích thước: ').bold = True
    p.add_run('D=3/8')  # ERROR: Sai ký hiệu đơn vị và thiếu đơn vị đo
    
    p = doc.add_paragraph()
    p.add_run('Áp suất hoạt động: ').bold = True
    p.add_run('Bao dải 0.3 - 0.95 Mpa (3 - 9.5 bar)')  # ERROR: Sai dấu thập phân mà dấu . thay vì , và Mpa thay vì MPa
    
    p = doc.add_paragraph()
    p.add_run('Dung tích: ').bold = True
    p.add_run('100 150 ml')  # ERROR: Thiếu ký hiệu phạm vi
    
    # Table
    doc.add_heading('III. CHI TIẾT MẶT HÀNG', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Mã hàng'
    hdr_cells[1].text = 'Tên hàng'
    hdr_cells[2].text = 'Số lượng'
    
    # Data
    rows_data = [
        ('001', 'Sản phẩm A', '100'),
        ('002', 'Sản phẩm B', '50'),
        ('003', 'Sản phẩm C', '75'),
    ]
    
    for i, (code, name, qty) in enumerate(rows_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = code
        row_cells[1].text = name
        row_cells[2].text = qty
    
    # Terms
    doc.add_heading('III. ĐIỀU KIỆN THANH TOÁN', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Giá trị hợp đồng: ').bold = True
    p.add_run('1.000.000 VND')
    
    p = doc.add_paragraph()
    p.add_run('Thanh toán: ').bold = True
    p.add_run('Chuyển khoản')
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.rows[0].cells[0].text = 'Bên mua (A)\nChữ ký: ___________'
    sig_table.rows[0].cells[1].text = 'Bên bán (B)\nChữ ký: ___________'
    
    return doc


def create_sample_regulation_document():
    """Tạo tài liệu quy định mẫu"""
    doc = Document()
    
    title = doc.add_heading('QUY ĐỊNH THÀNH LẬP HỢP ĐỒNG VÀ TRÌNH BÀY ĐƠN VỊ ĐO', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading('Phụ lục II - Thiết lập bội thập phân, ước thập phân của đơn vị đo', level=2)
    doc.add_paragraph('Tiền tố "mega" được ký hiệu là "M", đơn vị "pascal" được ký hiệu là "Pa". Do đó, "Mpa" phải được viết thành "MPa" (ký hiệu đúng là "MPa") để thể hiện megapascal.')
    doc.add_paragraph('Các đơn vị đo tiêu chuẩn bao gồm: mm, cm, m, km, mm2, cm2, m2, bar, Pa, MPa, K, °C, V, A, W, Hz, Ω')
    
    doc.add_heading('Phụ lục III - ĐƠN VỊ ĐO THEO THÔNG LỆ QUỐC TẾ', level=2)
    doc.add_paragraph('Chỉ được sử dụng các đơn vị đo trong danh sách tiêu chuẩn được phê duyệt.')
    doc.add_paragraph('Danh sách các đơn vị được phép: mm, cm, m, bar, Pa, MPa (không được dùng "D", "mpa", "Mpa")')
    
    doc.add_heading('Phụ lục V - TRÌNH BÀY ĐƠN VỊ ĐO PHÁP ĐỊNH', level=2)
    doc.add_paragraph('Khi thể hiện giá trị đại lượng đo, ký hiệu đơn vị đo phải đặt sau trị số, giữa hai thành phần này phải cách nhau một dấu cách.')
    doc.add_paragraph('Giá trị đại lượng đo phải kèm theo đơn vị đo cụ thể (ví dụ: "3/8 mm" hoặc "0,375 mm" không phải "3/8" hoặc "D=3/8").')
    doc.add_paragraph('Biểu thị dấu thập phân của giá trị đại lượng đo phải sử dụng dấu phẩy (,) không sử dụng dấu chấm (.). Ví dụ: "0,95 MPa" không phải "0.95 Mpa".')
    
    doc.add_heading('1. Yêu cầu chung', level=2)
    doc.add_paragraph('Tất cả hợp đồng phải được lập theo mẫu chuẩn được phê duyệt')
    doc.add_paragraph('Công ty phải được viết dưới dạng: "Công ty tnhh [Tên]" hoặc "Công ty cổ phần [Tên]"')
    doc.add_paragraph('Ngày tháng năm phải được viết dưới dạng: DD/MM/YYYY')
    
    doc.add_heading('2. Thông tin bên giao dịch', level=2)
    doc.add_paragraph('Phải ghi đầy đủ tên và chức vụ của đại diện pháp luật')
    doc.add_paragraph('Không được để trống bất kỳ trường nào')
    
    doc.add_heading('3. Tiền tệ', level=2)
    doc.add_paragraph('Sử dụng VND (Việt Nam Đồng) hoặc USD (Đô la Mỹ)')
    doc.add_paragraph('Phải nhất quán trên toàn bộ hợp đồng')
    
    return doc


def create_sample_reference_document():
    """Tạo tài liệu sở cứ mẫu"""
    doc = Document()
    
    title = doc.add_heading('MẪU HỢP ĐỒNG MUA BÁN HÀNG HÓA CHUẨN', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph('Đây là mẫu hợp đồng theo chuẩn của công ty')
    
    doc.add_heading('Phần 1: Tiêu đề', level=2)
    doc.add_paragraph('Tiêu đề: HỢP ĐỒNG MUA BÁN HÀNG HÓA')
    
    doc.add_heading('Phần 2: Bên giao dịch', level=2)
    p = doc.add_paragraph()
    p.add_run('Bên mua: ').bold = True
    p.add_run('Công ty tnhh/cổ phần [Tên công ty]')
    
    p = doc.add_paragraph()
    p.add_run('Bên bán: ').bold = True
    p.add_run('Công ty tnhh/cổ phần [Tên công ty]')
    
    return doc


def create_sample_documents(output_dir='samples'):
    """Tạo tất cả tài liệu mẫu"""
    os.makedirs(output_dir, exist_ok=True)
    
    print("Creating sample documents...")
    
    # Main document
    main_doc = create_sample_main_document()
    main_path = os.path.join(output_dir, 'sample_main.docx')
    main_doc.save(main_path)
    print(f"✓ Created: {main_path}")
    
    # Regulation
    reg_doc = create_sample_regulation_document()
    reg_path = os.path.join(output_dir, 'sample_regulation.docx')
    reg_doc.save(reg_path)
    print(f"✓ Created: {reg_path}")
    
    # Reference
    ref_doc = create_sample_reference_document()
    ref_path = os.path.join(output_dir, 'sample_reference.docx')
    ref_doc.save(ref_path)
    print(f"✓ Created: {ref_path}")
    
    print(f"\nAll sample documents created in: {output_dir}/")
    print("\nYou can now upload these files to test the application!")


if __name__ == '__main__':
    create_sample_documents()
